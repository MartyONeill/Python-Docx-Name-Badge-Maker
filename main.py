from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy

docx_file="template.docx"
doc = Document(docx_file)

df = pd.read_csv("names.csv")

#print(df)

def update_cell(name, institution, row_index):

    if institution is numpy.NAN:
        institution = ""

    counter = 0
    for table_idx, table in enumerate(doc.tables):

        count = 0

        for row_idx, row in enumerate(table.rows):

            for cell_idx, cell in enumerate(row.cells):

                count += 1

                if count is row_index:

                    if "{{institution}}" in cell.text:
                        cell.text = cell.text.replace("{{institution}}", institution)
                    if "{{name}}" in cell.text:
                        cell.text = cell.text.replace("{{name}}", name)
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

for index, item in df.iterrows():
    #print(item["person"], item["institution"], index+1)
    update_cell(item["person"], item["institution"], index+1)

doc.save(docx_file)